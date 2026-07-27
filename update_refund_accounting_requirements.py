from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path("Đặc tả trường dữ liệu sinh bút toán thanh toán - v3.docx")
TARGET = Path("Đặc tả trường dữ liệu sinh bút toán thanh toán - v4.docx")
BLACK = "000000"
WHITE = "FFFFFF"

CATALOG_ROWS = [
    ["refund_amount", "Số tiền hoàn ứng lần này", "esdHTKTpaymentVendor.refund.amount", "Có"],
    ["payable_account", "TK Phải trả NCC", "esdHTKTvendorSite.credit.account", "Có"],
    ["prepayment_account", "TK Tạm ứng", "esdHTKTvendorSite.debit.account", "Có"],
    ["vendor_site_id", "Mã site NCC để lấy tài khoản", "esdHTKTpaymentVendor.vendor.site.id", "Có"],
]

CASE_ROWS = [
    ["Số tiền hoàn ứng lần này", "esdHTKTpaymentVendor", "refund.amount", "Có", "G5"],
    ["TK Phải trả NCC", "esdHTKTvendorSite", "credit.account", "Có", "G5"],
    ["TK Tạm ứng", "esdHTKTvendorSite", "debit.account", "Có", "G5"],
    ["Mã site NCC để lấy tài khoản", "esdHTKTpaymentVendor", "vendor.site.id", "Có", "G5"],
]


def set_fill(cell, fill=WHITE):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_text(cell, text, size=7.9):
    cell.text = str(text)
    set_fill(cell, WHITE)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(BLACK)
            run.bold = False


def replace_text_in_paragraph(paragraph, replacements):
    old = paragraph.text
    new = old
    for source, target in replacements:
        new = new.replace(source, target)
    if new == old:
        return False
    paragraph.text = new
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    return True


doc = Document(SOURCE)
catalog_updated = 0
case_tables_updated = 0

# Thay 4 dòng G5 trong danh mục trường dùng chung.
for table in doc.tables:
    for row_index in range(1, len(table.rows) - 3):
        first = table.rows[row_index].cells[0].text.strip()
        if first != "prepayment_apply[]":
            continue
        if len(table.rows[row_index].cells) != 4:
            continue
        for offset, values in enumerate(CATALOG_ROWS):
            row = table.rows[row_index + offset]
            for column, value in enumerate(values):
                set_text(row.cells[column], value, 8.0)
        catalog_updated += 1

# Thay 4 dòng G5 trong từng case, giữ nguyên vị trí và cấu trúc bảng.
for table in doc.tables:
    g5_indexes = []
    for row_index, row in enumerate(table.rows):
        if len(row.cells) >= 5 and row.cells[-1].text.strip() == "G5":
            g5_indexes.append(row_index)
    if len(g5_indexes) != 4:
        continue
    for row_index, values in zip(g5_indexes, CASE_ROWS):
        row = table.rows[row_index]
        for column, value in enumerate(values):
            set_text(row.cells[column], value, 7.7)
    case_tables_updated += 1

replacements = [
    (
        "3.1. Bảng chi tiết hoàn ứng — GIẢ THIẾT: esdHTKTpaymentPrepaymentApply",
        "3.1. Chi tiết khoản tạm ứng nguồn — không bắt buộc để sinh bút toán kế toán",
    ),
    (
        "Tổng apply theo khoản = (3) = (1)",
        "Số tiền hoàn ứng lần này (3) = (1)",
    ),
    (
        "tổng apply=(3)",
        "sử dụng refund.amount=(3)",
    ),
]

text_updates = 0
for paragraph in doc.paragraphs:
    if replace_text_in_paragraph(paragraph, replacements):
        text_updates += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_text_in_paragraph(paragraph, replacements):
                    text_updates += 1

# Bổ sung ghi chú phạm vi ngay trước bảng giả thiết tích hợp nếu tìm thấy heading 3.1.
for paragraph in doc.paragraphs:
    if paragraph.text.startswith(
        "3.1. Chi tiết khoản tạm ứng nguồn — không bắt buộc"
    ):
        paragraph.text += (
            ". Phần dưới chỉ phục vụ tích hợp/đối soát về sau; "
            "không phải điều kiện chặn sinh TT-BK-04 và TT-BK-05."
        )
        break

# Đồng bộ table indent với cell margin 120 DXA của tài liệu người dùng.
for table in doc.tables:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

doc.core_properties.comments = (
    "Cập nhật G5: đủ dữ liệu sinh bút toán hoàn ứng từ "
    "paymentVendor.refund.amount và tài khoản debit/credit của vendorSite. "
    "Chi tiết giao dịch nguồn chỉ phục vụ tích hợp/đối soát."
)
doc.save(TARGET)
print(
    f"catalog_updated={catalog_updated}; "
    f"case_tables_updated={case_tables_updated}; text_updates={text_updates}"
)
