from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path("Đặc tả trường dữ liệu sinh bút toán thanh toán - v2.docx")

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EFF6FB"
PALE_YELLOW = "FFF4CE"
PALE_GREEN = "E2F0D9"
PALE_RED = "FCE4D6"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=110):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, size=9.2, bold=False, color=BLACK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def write_cell(cell, text, bold=False, color=BLACK, size=8.4, align=None):
    is_assumption = "GIẢ THIẾT" in str(text).upper()
    if is_assumption:
        set_cell_shading(cell, PALE_YELLOW)
        color = "C00000"
        bold = True
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    parts = str(text).split("\n")
    for idx, part in enumerate(parts):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        style_run(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        write_cell(
            table.rows[0].cells[i],
            header,
            bold=True,
            color=WHITE,
            size=8.2,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if row_idx % 2:
                set_cell_shading(cells[i], "F8FBFD")
            write_cell(cells[i], value, size=font_size)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_callout(doc, label, text, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    style_run(r, size=9.2, bold=True, color=NAVY)
    r = p.add_run(text)
    style_run(r, size=9.2)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_together = True
        r = p.add_run("☐ " + item)
        style_run(r, size=9.2)


def add_case(doc, case):
    add_heading(doc, f"{case['code']} — {case['name']}", 2)
    add_table(
        doc,
        ["Điều kiện nhận diện", "Nhóm dữ liệu sử dụng", "API/OGL", "Kết quả cần sinh"],
        [[case["condition"], case["groups"], case["api"], case["result"]]],
        [2600, 1900, 1700, 3160],
        8.4,
    )
    case_fields = []
    for group_code in [x.strip() for x in case["groups"].split("+")]:
        for logical_code, vietnamese_name, db_mapping, required in CASE_FIELD_GROUPS[group_code]:
            mapping = str(db_mapping)
            if mapping.startswith("CẦN ") or mapping.startswith("Chưa "):
                db_table = ""
                db_field = ""
            elif "." in mapping:
                db_table, db_field = mapping.split(".", 1)
            else:
                db_table = mapping
                db_field = "Cần xác nhận trường cụ thể"
            case_fields.append(
                [
                    vietnamese_name,
                    db_table,
                    db_field,
                    required,
                    group_code,
                ]
            )
    add_table(
        doc,
        ["Tên trường nghiệp vụ (tiếng Việt)", "Bảng DB", "Trường DB", "Bắt buộc", "Nhóm"],
        case_fields,
        [2920, 1900, 2780, 940, 820],
        7.7,
    )
    add_callout(
        doc,
        "Điều kiện hoàn tất case",
        case["gate"] + "  Trạng thái xác nhận: ☐ Chưa xác nhận  ☐ Đã xác nhận",
        PALE_GREEN,
    )


common = [
    ("payment_id", "Mã phiếu thanh toán/referenceId", "esdHTKTpayment.id", "Có"),
    ("vendor_id", "Mã NCC nội bộ", "esdHTKTpaymentVendor.vendor.id", "Có"),
    ("vendor_number", "Mã NCC/MST/CCCD dùng OGL", "esdHTKTvendor.vendor.number", "Có"),
    ("vendor_type", "Tổ chức hoặc Cá nhân", "esdHTKTvendor.vendor.type", "Có"),
    ("vendor_site_id", "Site NCC được chọn", "esdHTKTpaymentVendor.vendor.site.id", "Có"),
    ("vendor_site_code", "Site code OGL", "esdHTKTvendorSite.ogl.site.code", "Có"),
    ("ogl_entity", "Đơn vị/entity gửi OGL", "esdHTKTvendorSite.ogl.entity", "Có"),
    ("payable_account", "Tài khoản phải trả NCC", "esdHTKTvendorSite.credit.account", "Có"),
    ("currency", "Loại tiền 3 ký tự", "esdHTKTpaymentVendor.currency", "Có"),
    ("payment_method", "Chuyển khoản/Tiền mặt", "esdHTKTpaymentVendor.payment.method", "Có"),
    ("accepted_invoice_amount", "Giá trị hóa đơn chấp nhận (1)", "GIẢ THIẾT TỪ CODE SAVE: esdHTKTpaymentVendor.approved.invoice.amount", "Theo case"),
    ("payment_requested_amount", "Số tiền đề nghị thanh toán (2)", "GIẢ THIẾT TỪ CODE SAVE: esdHTKTpaymentVendor.amount", "Theo case"),
    ("refund_amount", "Tổng hoàn ứng lần này (3)", "GIẢ THIẾT TỪ CODE SAVE: esdHTKTpaymentVendor.refund.amount", "Theo case"),
    ("transaction_description", "Nội dung giao dịch", "esdHTKTpaymentVendor.transaction.des", "Có"),
    ("maker/checker", "Người tạo/người duyệt cuối", "payment.created.by / user.approver.final", "Có"),
    ("contract_id", "Mã HĐ/KMS", "esdHTKTpayment.contract.id", "Không"),
]

invoice_fields = [
    ("invoice_links[]", "Danh sách hóa đơn của đúng NCC", "esdHTKTpaymentInvoice + esdHTKTinvoice", "Có"),
    ("cost_divisions[]", "Các dòng chi phí trước thuế", "esdHTKTpaymentCostDivision theo payment+vendor", "Có"),
    ("cost_account_segments", "Đủ segment1…segment7 OGL", "GIẢ THIẾT: esdHTKTpaymentCostDivision.account.number lưu full 7 segment", "Có"),
    ("tax_amount", "Thuế được hạch toán/khấu trừ", "paymentInvoice.deduction.amount", "Nếu có thuế"),
    ("deduction_type/rate", "Loại và tỷ lệ khấu trừ", "paymentInvoice.deduction.type/rate", "Nếu có thuế"),
]

personal_tax = [
    ("pit_amount", "Thuế TNCN cần khấu trừ", "GIẢ THIẾT: esdHTKTpaymentPersonalTax.amount", "Có"),
    ("pit_rate", "Tỷ lệ thuế TNCN", "GIẢ THIẾT: esdHTKTpaymentPersonalTax.rate", "Có"),
    ("pit_expense_account", "TK chi phí thuế TNCN", "GIẢ THIẾT: esdHTKTpaymentPersonalTax.expense.account", "Có"),
    ("pit_payable_account", "TK thuế TNCN phải nộp", "GIẢ THIẾT: esdHTKTpaymentPersonalTax.payable.account", "Có"),
]

payment_fields = [
    ("beneficiary_account", "Tài khoản thụ hưởng", "paymentVendor.beneficiary.account", "Nếu chuyển khoản"),
    ("beneficiary_name", "Tên thụ hưởng", "paymentVendor.beneficiary.name", "Nếu chuyển khoản"),
    ("beneficiary_bank_code", "Mã ngân hàng/chi nhánh thụ hưởng", "GIẢ THIẾT NGỮ NGHĨA: esdHTKTpaymentVendor.beneficiary.bank", "Nếu chuyển khoản"),
    ("cash_account", "Tài khoản quỹ tiền mặt", "GIẢ THIẾT: esdHTKTcashAccountConfig.account.number", "Nếu tiền mặt"),
]

refund_fields = [
    ("prepayment_apply[]", "Danh sách khoản tạm ứng được hoàn", "GIẢ THIẾT: esdHTKTpaymentPrepaymentApply.id", "Có"),
    ("prepayment_transaction_id", "Mã giao dịch/mã invoice OGL tạm ứng", "GIẢ THIẾT: esdHTKTpaymentPrepaymentApply.prepayment.transaction.id", "Có"),
    ("prepayment_apply_amount", "Số hoàn ứng theo từng khoản", "GIẢ THIẾT: esdHTKTpaymentPrepaymentApply.amount", "Có"),
    ("prepayment_account", "TK tạm ứng của khoản gốc", "GIẢ THIẾT: esdHTKTpaymentPrepaymentApply.prepayment.account", "Có"),
]

suspended_fields = [
    ("payable_apply[]", "Danh sách khoản phải trả cũ", "GIẢ THIẾT: esdHTKTpaymentPayableApply.id", "Có"),
    ("source_invoice_number", "Số YCTT/invoiceNumber OGL của phiếu gốc", "GIẢ THIẾT: esdHTKTpaymentPayableApply.source.invoice.number", "Có"),
    ("payable_apply_amount", "Số thanh toán cho từng khoản cũ", "GIẢ THIẾT: esdHTKTpaymentPayableApply.amount", "Có"),
]

CASE_FIELD_GROUPS = {
    "G0": common,
    "G1": invoice_fields[:3],
    "G2": invoice_fields[3:],
    "G3": payment_fields,
    "G4": personal_tax,
    "G5": refund_fields,
    "G6": suspended_fields,
}


def fields_for(*groups):
    result = []
    seen = set()
    for group in groups:
        for row in group:
            if row[0] not in seen:
                result.append(row)
                seen.add(row[0])
    return result


cases = [
    dict(code="TT-01", name="Thanh toán toàn bộ, không thuế", condition="(3)=0; (1)=(2); không thuế", groups="G0 + G1 + G3", api="/ap/create-invoice\nSTANDARD + PAYMENT", result="Nợ Chi phí (1)\nCó TK chi tiền (1)", gate="Tổng phân bổ = (1); (2)=(1); không có dòng thuế/hoàn ứng/khoản treo."),
    dict(code="TT-02", name="Thanh toán toàn bộ, có thuế — Cá nhân", condition="(3)=0; (1)=(2); có thuế; Cá nhân", groups="G0 + G1 + G3 + G4", api="/ap/create-invoice\n+ /general-ledger/interface", result="AP: Nợ Chi phí, Có TK chi tiền\nGL: Nợ CP thuế TNCN, Có Thuế TNCN", gate="Xác định số thực chi sau TNCN và hai tài khoản GL; AP+GL cân đối."),
    dict(code="TT-03", name="Thanh toán toàn bộ, có thuế — NCC", condition="(3)=0; (1)=(2); có thuế; Tổ chức", groups="G0 + G1 + G2 + G3", api="/ap/create-invoice\nSTANDARD + PAYMENT", result="Nợ Chi phí trước thuế\nNợ Thuế GTGT\nCó TK chi tiền (2)", gate="Chi phí + thuế = (1) và (2)=(1)."),
    dict(code="TT-04", name="Thanh toán một phần, giữ phải trả — Không thuế", condition="(3)=0; (1)>(2); không thuế", groups="G0 + G1 + G3", api="/ap/create-invoice\nSTANDARD + PAYMENT", result="Nợ Chi phí (1)\nCó TK chi tiền (2)\nCó Phải trả (1)-(2)", gate="Số phải trả còn lại = (1)-(2) > 0."),
    dict(code="TT-05", name="Thanh toán một phần, giữ phải trả — Cá nhân có thuế", condition="(3)=0; (1)>(2); có thuế; Cá nhân", groups="G0 + G1 + G3 + G4", api="/ap/create-invoice\n+ /general-ledger/interface", result="AP: Chi phí/chi tiền/phải trả\nGL: Chi phí thuế TNCN/Thuế TNCN", gate="Chốt công thức phân bổ (2) giữa thực chi, thuế TNCN và phải trả."),
    dict(code="TT-06", name="Thanh toán một phần, giữ phải trả — NCC có thuế", condition="(3)=0; (1)>(2); có thuế; Tổ chức", groups="G0 + G1 + G2 + G3", api="/ap/create-invoice\nSTANDARD + PAYMENT", result="Nợ Chi phí trước thuế\nNợ Thuế GTGT\nCó TK chi tiền (2)\nCó Phải trả (1)-(2)", gate="Chi phí + thuế = (1); tiền + phải trả = (1)."),
    dict(code="TT-07", name="Hoàn ứng toàn bộ, không thanh toán thêm", condition="(3)=(1); (2)=0", groups="G0 + G1 + G5", api="/ap/create-invoice\nSTANDARD + APPLY_PREPAYMENT", result="Nợ Chi phí (3)\nCó TK Tạm ứng (3)", gate="Tổng apply theo khoản = (3) = (1); không sinh chuyển tiền."),
    dict(code="TT-08", name="Hoàn một phần, thanh toán phần còn lại — Không thuế", condition="(3)>0; (2)>0; (2)+(3)=(1); không thuế", groups="G0 + G1 + G3 + G5", api="/ap/create-invoice\nSTANDARD + APPLY_PREPAYMENT + PAYMENT", result="Nợ Chi phí (1)\nCó Tạm ứng (3)\nCó TK chi tiền (2)", gate="(2)+(3)=(1); tổng apply=(3); không còn phải trả."),
    dict(code="TT-09", name="Hoàn một phần, thanh toán phần còn lại — NCC có thuế", condition="(2)+(3)=(1); có thuế; Tổ chức", groups="G0 + G1 + G2 + G3 + G5", api="/ap/create-invoice\nSTANDARD + APPLY_PREPAYMENT + PAYMENT", result="Nợ Chi phí trước thuế\nNợ Thuế GTGT\nCó Tạm ứng (3)\nCó TK chi tiền (2)", gate="Chi phí + thuế = (1); (2)+(3)=(1)."),
    dict(code="TT-10", name="Hoàn một phần, thanh toán phần còn lại — Cá nhân có thuế", condition="(2)+(3)=(1); có thuế; Cá nhân", groups="G0 + G1 + G3 + G4 + G5", api="/ap/create-invoice\n+ /general-ledger/interface", result="AP: Chi phí/Tạm ứng/Chi tiền\nGL: Chi phí thuế TNCN/Thuế TNCN", gate="Chốt quan hệ (1),(2),(3) với thuế TNCN; AP+GL cân đối."),
    dict(code="TT-11", name="Hoàn một phần, giữ phải trả — Không thuế", condition="(3)>0; (2)=0; (1)>(3)", groups="G0 + G1 + G5", api="/ap/create-invoice\nSTANDARD + APPLY_PREPAYMENT", result="Nợ Chi phí (1)\nCó Tạm ứng (3)\nCó Phải trả (1)-(3)", gate="Số phải trả còn lại = (1)-(3) > 0."),
    dict(code="TT-12", name="Hoàn một phần, giữ phải trả — NCC có thuế", condition="(3)>0; (2)=0; (1)>(3); Tổ chức", groups="G0 + G1 + G2 + G5", api="/ap/create-invoice\nSTANDARD + APPLY_PREPAYMENT", result="Nợ Chi phí trước thuế\nNợ Thuế GTGT\nCó Tạm ứng (3)\nCó Phải trả (1)-(3)", gate="Chi phí + thuế = (1); tạm ứng + phải trả = (1)."),
    dict(code="TT-13", name="Hoàn một phần, giữ phải trả — Cá nhân có thuế", condition="(3)>0; (2)=0; (1)>(3); Cá nhân", groups="G0 + G1 + G4 + G5", api="/ap/create-invoice\n+ /general-ledger/interface", result="AP: Chi phí/Tạm ứng/Phải trả\nGL: Chi phí thuế TNCN/Thuế TNCN", gate="Chốt số phải trả sau TNCN; tổng AP+GL cân đối."),
    dict(code="TT-14", name="Hoàn một phần, thanh toán một phần, giữ phải trả — Không thuế", condition="(2)>0; (3)>0; (2)+(3)<(1)", groups="G0 + G1 + G3 + G5", api="/ap/create-invoice\nSTANDARD + APPLY_PREPAYMENT + PAYMENT", result="Nợ Chi phí (1)\nCó Tạm ứng (3)\nCó TK chi tiền (2)\nCó Phải trả (1)-(2)-(3)", gate="Phải trả = (1)-(2)-(3) > 0."),
    dict(code="TT-15", name="Hoàn một phần, thanh toán một phần, giữ phải trả — NCC có thuế", condition="(2)+(3)<(1); có thuế; Tổ chức", groups="G0 + G1 + G2 + G3 + G5", api="/ap/create-invoice\nSTANDARD + APPLY_PREPAYMENT + PAYMENT", result="Nợ Chi phí trước thuế\nNợ Thuế GTGT\nCó Tạm ứng/Chi tiền/Phải trả", gate="Chi phí + thuế = tạm ứng + tiền + phải trả = (1)."),
    dict(code="TT-16", name="Hoàn một phần, thanh toán một phần, giữ phải trả — Cá nhân", condition="(2)+(3)<(1); có thuế; Cá nhân", groups="G0 + G1 + G3 + G4 + G5", api="/ap/create-invoice\n+ /general-ledger/interface", result="AP: Chi phí/Tạm ứng/Chi tiền/Phải trả\nGL: Chi phí thuế TNCN/Thuế TNCN", gate="Chốt công thức phải trả sau TNCN; AP+GL cân đối."),
    dict(code="TT-17", name="Chỉ thanh toán khoản phải trả của YCTT trước", condition="(1)=0; (2)>0; không ghi nhận chi phí mới", groups="G0 + G3 + G6", api="/ap/create-payment", result="Nợ TK Phải trả (2)\nCó TK chi tiền (2)", gate="Mỗi khoản có invoiceNumber/transactionId OGL gốc; tổng apply=(2); không có cost division."),
]


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(9.2)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.12

for name, size, before, after, color in [
    ("Heading 1", 15, 14, 7, NAVY),
    ("Heading 2", 11.5, 10, 5, BLUE),
    ("Heading 3", 10, 8, 4, NAVY),
]:
    s = styles[name]
    s.font.name = "Arial"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("ĐẶC TẢ ĐẦU VÀO BÚT TOÁN THANH TOÁN  |  DRAFT")
style_run(hr, size=7.8, bold=True, color=GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Tài liệu làm rõ dữ liệu trước khi triển khai sinh bút toán")
style_run(fr, size=7.5, color=GRAY)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("ĐẶC TẢ TRƯỜNG DỮ LIỆU")
style_run(r, size=22, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run("Sinh bút toán kế toán cho nghiệp vụ Thanh toán")
style_run(r, size=14, bold=True, color=BLUE)

add_table(
    doc,
    ["Thuộc tính", "Giá trị"],
    [
        ["Trạng thái", "DRAFT — Chưa được dùng để lập trình sinh bút toán"],
        ["Phạm vi", "17 case TT-01 đến TT-17; AP, GL và Core Banking"],
        ["Mục đích", "Hoàn thiện nguồn dữ liệu, công thức, mapping và điều kiện kiểm tra trước khi code"],
        ["Người xác nhận nghiệp vụ", "........................................................"],
        ["Người xác nhận DB/tích hợp", "........................................................"],
        ["Ngày chốt phiên bản", "........ / ........ / ........"],
    ],
    [1900, 7460],
    8.7,
)

add_callout(
    doc,
    "Nguyên tắc sử dụng",
    "Chỉ bắt đầu sửa/sinh code bút toán khi toàn bộ trường P0, công thức số tiền và nguồn DB của từng case đã được đánh dấu “Đã xác nhận”.",
    PALE_RED,
)

add_heading(doc, "1. Quy ước và điều kiện nền", 1)
add_callout(
    doc,
    "Phân loại bằng chứng DB",
    "Ảnh bảng vừa được cung cấp thuộc phần Tạm ứng và chỉ dùng để tham khảo; không dùng ảnh đó để xác nhận bảng Thanh toán. Ba mapping (1), (2), (3) dưới đây được suy ra từ đoạn code save vào $L_file và vẫn là GIẢ THIẾT cho đến khi có ảnh Fields hoặc descriptor xác nhận $L_file là esdHTKTpaymentVendor.",
    PALE_YELLOW,
)
add_table(
    doc,
    ["Ký hiệu", "Định nghĩa phải chốt", "Trường DB đề xuất", "Trạng thái"],
    [
        ["(1)", "Giá trị hóa đơn chấp nhận theo từng NCC", "GIẢ THIẾT TỪ CODE SAVE: esdHTKTpaymentVendor.approved.invoice.amount", "☐"],
        ["(2)", "Số tiền đề nghị thanh toán theo từng NCC; chưa tự động đồng nghĩa với số thực chi", "GIẢ THIẾT TỪ CODE SAVE: esdHTKTpaymentVendor.amount", "☐"],
        ["(3)", "Tổng số tiền hoàn ứng lần này theo từng NCC", "GIẢ THIẾT TỪ CODE SAVE: esdHTKTpaymentVendor.refund.amount", "☐"],
        ["Thuế GTGT", "Số thuế được hạch toán theo loại/tỷ lệ khấu trừ", "PaymentInvoice.deduction_amount", "☐"],
        ["Thuế TNCN", "Số thuế khấu trừ đối với NCC cá nhân", "Cần bổ sung theo NCC/hóa đơn", "☐"],
        ["Phải trả còn lại", "Phần nghĩa vụ chưa được hoàn ứng/chưa chi trả", "(1) − (2) − (3), có điều chỉnh TNCN nếu áp dụng", "☐"],
    ],
    [900, 3630, 3530, 1300],
    8.2,
)

add_heading(doc, "2. Danh mục trường dùng chung", 1)
add_callout(
    doc,
    "Quy tắc chống lặp",
    "Mỗi trường chỉ được định nghĩa tại một nhóm G0–G6. Phần case chỉ tham chiếu mã nhóm, không chép lại danh sách trường.",
    PALE_BLUE,
)
field_groups = [
    ("G0", "Thông tin nền phiếu/NCC", CASE_FIELD_GROUPS["G0"]),
    ("G1", "Hóa đơn và phân bổ chi phí", CASE_FIELD_GROUPS["G1"]),
    ("G2", "Thuế GTGT", CASE_FIELD_GROUPS["G2"]),
    ("G3", "Chi trả/Chuyển tiền", CASE_FIELD_GROUPS["G3"]),
    ("G4", "Thuế TNCN cá nhân", CASE_FIELD_GROUPS["G4"]),
    ("G5", "Hoàn ứng", CASE_FIELD_GROUPS["G5"]),
    ("G6", "Trả khoản phải trả cũ", CASE_FIELD_GROUPS["G6"]),
]
for group_code, group_name, group_rows in field_groups:
    add_heading(doc, f"{group_code} — {group_name}", 2)
    displayed_group_rows = []
    for logical_code, vietnamese_name, db_mapping, required in group_rows:
        displayed_mapping = str(db_mapping)
        if displayed_mapping.startswith("CẦN ") or displayed_mapping.startswith("Chưa "):
            displayed_mapping = ""
        displayed_group_rows.append(
            [logical_code, vietnamese_name, displayed_mapping, required]
        )
    add_table(
        doc,
        ["Mã trường", "Mục đích", "Nguồn DB dự kiến", "Bắt buộc"],
        displayed_group_rows,
        [2180, 3200, 2860, 1120],
        8.0,
    )

add_heading(doc, "3. Cấu trúc DB cần hoàn thiện", 1)
add_callout(
    doc,
    "GIẢ THIẾT THIẾT KẾ DB",
    "Toàn bộ tên bảng/trường trong mục 3 là phương án đề xuất do thiết kế hiện tại chưa cung cấp đủ thông tin. Cần được đội DB xác nhận hoặc thay thế.",
    PALE_YELLOW,
)
add_heading(doc, "3.1. Bảng chi tiết hoàn ứng — GIẢ THIẾT: esdHTKTpaymentPrepaymentApply", 2)
add_table(
    doc,
    ["Trường", "Kiểu gợi ý", "Mục đích"],
    [
        ["GIẢ THIẾT: id", "varchar PK", "Khóa dòng áp khoản tạm ứng"],
        ["GIẢ THIẾT: payment.id, vendor.id", "varchar FK", "Xác định phiếu và NCC"],
        ["GIẢ THIẾT: prepayment.id", "varchar FK", "Mã phiếu tạm ứng nguồn"],
        ["GIẢ THIẾT: prepayment.transaction.id", "varchar", "Mã giao dịch/invoiceNumber do OGL trả về"],
        ["GIẢ THIẾT: amount", "decimal", "Số hoàn ứng lần này của khoản"],
        ["GIẢ THIẾT: currency", "varchar(3)", "Loại tiền"],
        ["GIẢ THIẾT: prepayment.account.number", "varchar", "Tài khoản tạm ứng nguồn"],
        ["GIẢ THIẾT: remaining.amount.before", "decimal", "Số còn có thể hoàn trước giao dịch"],
        ["GIẢ THIẾT: status", "varchar", "Dự kiến/Đang xử lý/Thành công/Lỗi"],
    ],
    [2600, 1700, 5060],
    8.3,
)

add_heading(doc, "3.2. Bảng chi tiết trả khoản phải trả cũ — GIẢ THIẾT: esdHTKTpaymentPayableApply", 2)
add_table(
    doc,
    ["Trường", "Kiểu gợi ý", "Mục đích"],
    [
        ["GIẢ THIẾT: id", "varchar PK", "Khóa dòng thanh toán khoản cũ"],
        ["GIẢ THIẾT: payment.id, vendor.id", "varchar FK", "Phiếu hiện tại và NCC"],
        ["GIẢ THIẾT: source.payment.id", "varchar FK", "Phiếu thanh toán nguồn"],
        ["GIẢ THIẾT: source.invoice.number", "varchar", "invoiceNumber/YCTT OGL của phiếu nguồn"],
        ["GIẢ THIẾT: source.transaction.id", "varchar", "Mã giao dịch OGL nguồn"],
        ["GIẢ THIẾT: amount", "decimal", "Số trả lần này"],
        ["GIẢ THIẾT: currency", "varchar(3)", "Loại tiền"],
        ["GIẢ THIẾT: payable.account.number", "varchar", "Tài khoản phải trả nguồn"],
        ["GIẢ THIẾT: remaining.amount.before", "decimal", "Số còn phải trả trước giao dịch"],
        ["GIẢ THIẾT: status", "varchar", "Dự kiến/Đang xử lý/Thành công/Lỗi"],
    ],
    [2600, 1700, 5060],
    8.3,
)

add_heading(doc, "3.3. Trường bổ sung cho PaymentEntry/AccountingInformation", 2)
add_table(
    doc,
    ["Nhóm", "Trường cần có", "Lý do"],
    [
        ["GIẢ THIẾT: PaymentEntry", "entry.code, source.type, source.id", "Nhận diện TT-BK-01…08 và ghép đúng dòng khi sinh lại"],
        ["GIẢ THIẾT: PaymentEntry", "segment1…segment7", "Tạo invoiceLineList/GL Line mà không suy diễn mơ hồ"],
        ["GIẢ THIẾT: PaymentEntry", "exchange.rate, accounted.amount", "Hạch toán ngoại tệ và số quy đổi"],
        ["GIẢ THIẾT: AccountingInformation", "endpoint, integration.type, idempotency.key", "Phân biệt AP Invoice/AP Payment/GL/Core và chống gửi trùng"],
        ["GIẢ THIẾT: AccountingInformation", "request.payload, response.payload dạng TEXT/CLOB", "Payload vượt quá varchar(60)"],
        ["GIẢ THIẾT: AccountingInformation", "transaction.id, invoice.number, payment.number, batch.name", "Đối soát, ApplyPrepayment, Case 17"],
        ["GIẢ THIẾT: AccountingInformation", "error.code, retry.count, timestamps", "Theo dõi lỗi và retry an toàn"],
    ],
    [2100, 3300, 3960],
    8.2,
)

add_heading(doc, "4. Trường bắt buộc theo từng case", 1)
add_callout(
    doc,
    "Cách hoàn thiện",
    "Các ô nền vàng/chữ đỏ có nhãn “GIẢ THIẾT” là phương án đề xuất, chưa được chứng minh từ thiết kế DB hiện có. Hãy xác nhận, sửa hoặc thay thế các giả thiết này bằng bảng/trường thực tế; sau đó kiểm tra công thức và đánh dấu case “Đã xác nhận”.",
    PALE_BLUE,
)
for case in cases:
    add_case(doc, case)

add_heading(doc, "5. Mapping tích hợp bắt buộc phải chốt", 1)
add_table(
    doc,
    ["API/nhóm", "Trường bắt buộc", "Nguồn dự kiến", "Trạng thái"],
    [
        ["/ap/create-invoice", "requestId, referenceId, vendorNumber, vendorSiteCode, entity", "AccountingInformation + Payment + Vendor/Site", "☐"],
        ["/ap/create-invoice", "invoiceType=STANDARD, amount=(1), amountPay=(2)", "PaymentVendor", "☐"],
        ["/ap/create-invoice", "invoiceLineList: amount + segment1…7", "PaymentCostDivision", "☐"],
        ["/ap/create-invoice", "applyList: invoiceNumber + amount", "PaymentPrepaymentApply", "☐"],
        ["/ap/create-invoice", "vatList: invoice id + deductionType", "PaymentInvoice", "☐"],
        ["/ap/create-payment", "invoiceNumber nguồn + amount", "PaymentPayableApply", "☐"],
        ["/general-ledger/interface", "AccountingDate, BranchCode, Source, Category, Line", "Payment + PaymentEntry GL", "☐"],
        ["Core Banking", "TK nguồn, TK đích, bank/branch, amount, currency, notes", "Cấu hình đơn vị + PaymentVendor", "☐"],
    ],
    [1850, 3180, 3130, 1200],
    8.0,
)

add_heading(doc, "6. Quy tắc kiểm tra trước khi sinh bút toán", 1)
add_checklist(
    doc,
    [
        "Phiếu/NCC có ít nhất một nguồn: hóa đơn mới, hoàn ứng hoặc trả khoản phải trả cũ.",
        "Tất cả số tiền dùng Decimal; thống nhất số chữ số thập phân và quy tắc làm tròn.",
        "Tổng Nợ = Tổng Có theo từng NCC, từng loại tiền và toàn phiếu.",
        "Tổng phân bổ chi phí + thuế = giá trị hóa đơn chấp nhận (1).",
        "Tổng chi tiết hoàn ứng = (3); mỗi khoản không vượt số còn có thể hoàn.",
        "Tổng chi tiết trả khoản cũ không vượt số phải trả còn lại của từng giao dịch nguồn.",
        "Không trừ hoàn ứng hai lần khi xác định số thực thanh toán.",
        "NCC cá nhân có thuế phải có đủ cặp bút toán GL TNCN.",
        "Mỗi dòng tự sinh có entry_code và source_id ổn định để giữ đúng chỉnh sửa khi sinh lại.",
        "Request tích hợp có idempotency key; retry không tạo giao dịch trùng.",
        "Không gọi Core Banking đối với phương thức Tiền mặt.",
        "Chỉ khóa/chốt entry sau khi phiếu qua giai đoạn KTTC được phép chỉnh sửa.",
    ],
)

add_heading(doc, "7. Các quyết định nghiệp vụ cần ký xác nhận", 1)
add_table(
    doc,
    ["STT", "Quyết định cần chốt", "Phương án được chọn/Ghi chú", "Xác nhận"],
    [
        ["1", "PaymentVendor.amount là (2) hay tổng trước khi trừ hoàn ứng?", "", "☐"],
        ["2", "Nguồn chính xác của (1) theo NCC?", "", "☐"],
        ["3", "Thuế GTGT dùng invoice.total.tax hay deduction.amount?", "", "☐"],
        ["4", "Công thức (2), (3), TNCN trong TT-02/05/10/13/16?", "", "☐"],
        ["5", "Tài khoản tiền mặt lấy từ cấu hình nào?", "", "☐"],
        ["6", "7 segment OGL lưu tách hay lưu full combination?", "", "☐"],
        ["7", "ApplyPrepayment cần transactionId hay invoiceNumber OGL?", "", "☐"],
        ["8", "Khoản treo có được sinh đồng thời với hóa đơn mới không?", "", "☐"],
        ["9", "Quy tắc ngoại tệ, tỷ giá và làm tròn?", "", "☐"],
        ["10", "Cách tạo requestId/idempotency key cho từng endpoint?", "", "☐"],
    ],
    [700, 4150, 3410, 1100],
    8.0,
)

add_heading(doc, "8. Biên bản sẵn sàng triển khai", 1)
add_callout(
    doc,
    "Điều kiện GO",
    "17/17 case đã xác nhận; không còn giả thiết chưa được phê duyệt; mapping OGL/Core hoàn tất; công thức được KTTC ký xác nhận; DB migration đã sẵn sàng.",
    PALE_GREEN,
)
add_table(
    doc,
    ["Hạng mục", "Kết quả", "Người xác nhận", "Ngày"],
    [
        ["Nghiệp vụ và công thức", "☐ Đạt  ☐ Chưa đạt", "", ""],
        ["Thiết kế DB", "☐ Đạt  ☐ Chưa đạt", "", ""],
        ["Mapping OGL", "☐ Đạt  ☐ Chưa đạt", "", ""],
        ["Mapping Core Banking", "☐ Đạt  ☐ Chưa đạt", "", ""],
        ["17 case và kiểm thử", "☐ Đạt  ☐ Chưa đạt", "", ""],
        ["Quyết định triển khai code", "☐ GO  ☐ NO-GO", "", ""],
    ],
    [3000, 2200, 2660, 1500],
    8.4,
)

doc.core_properties.title = "Đặc tả trường dữ liệu sinh bút toán thanh toán"
doc.core_properties.subject = "Checklist hoàn thiện dữ liệu trước khi triển khai 17 case bút toán"
doc.core_properties.author = "Codex"
doc.save(OUT)
print("document-created")
