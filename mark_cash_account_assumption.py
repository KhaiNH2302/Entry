from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path("Đặc tả trường dữ liệu sinh bút toán thanh toán - v2.docx")
TARGET = Path("Đặc tả trường dữ liệu sinh bút toán thanh toán - v3.docx")
LABEL = "Tài khoản quỹ tiền mặt"
YELLOW = "FFF4CE"
RED = "C00000"


def shade(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), YELLOW)


def mark_cell(cell, prefix=False):
    original = cell.text.strip()
    if prefix and not original.upper().startswith("GIẢ THIẾT"):
        original = "GIẢ THIẾT: " + original
    cell.text = original
    shade(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
            run.font.size = Pt(7.9)
            run.font.color.rgb = RGBColor.from_string(RED)
            run.bold = True


doc = Document(SOURCE)
updated = 0

for table in doc.tables:
    for row in table.rows:
        values = [cell.text.strip() for cell in row.cells]
        if not values or LABEL.lower() not in values[0].lower():
            continue

        # Danh mục G3: Mã trường | Tên tiếng Việt | Map bảng.trường | Bắt buộc
        if len(row.cells) == 4:
            mark_cell(row.cells[2], prefix=True)
            updated += 1
            continue

        # Bảng theo case: Tên tiếng Việt | Bảng DB | Trường DB | Bắt buộc | Nhóm
        if len(row.cells) >= 5:
            mark_cell(row.cells[1], prefix=True)
            mark_cell(row.cells[2], prefix=False)
            updated += 1

# Giữ hình học bảng nhất quán với cell margin 120 DXA của bản người dùng đã sửa.
for table in doc.tables:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

doc.core_properties.comments = (
    "Tài khoản quỹ tiền mặt được giữ ở trạng thái GIẢ THIẾT; "
    "các ô mapping được đánh dấu nền vàng, chữ đỏ."
)
doc.save(TARGET)
print(f"updated={updated}")
